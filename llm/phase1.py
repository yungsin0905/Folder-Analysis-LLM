import os
import io
import csv
import re
from typing import List
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI

import pandas as pd
from docx import Document
from PIL import Image
from pdf2image import convert_from_bytes
import pytesseract
from dateutil import parser
import base64
import uuid
import asyncio

# 存储每个任务的处理进度： {job_id: {"total": int, "done": int, "status": "processing"/"finished"/"cancelled", "result": bytes}}
JOBS = {}

# 存储被取消的任务 ID
CANCELLED_JOBS = set()

app = FastAPI()

# 跨域设置，允许前端页面访问 API
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ------------------------------------------------------------------
# 0. 首页
# ------------------------------------------------------------------
@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    html_path = os.path.join(os.path.dirname(__file__), "..", "index.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return f.read()

# ------------------------------------------------------------------
# 1. 环境与 LM Studio 配置
# ------------------------------------------------------------------
LM_STUDIO_URL = "http://127.0.0.1:1234/v1"

client = OpenAI(
    base_url=LM_STUDIO_URL, 
    api_key="lm-studio"
)

# 环境路径
POPPLER_PATH = r"C:\poppler-25.12.0\Library\bin"
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

DOC_TYPE_MAP = {
    "LO":        "Letter of Offer",
    "TA":        "Tenancy Agreement",
    "RENEWAL":   "Renewal of Tenancy",
    "HANDOVER":  "Handover Form",
    "DEED":      "Deed of Assignment",
    "COVER":     "Cover Letter",
    "DEFAULT":   "Correspondence"
}

# ------------------------------------------------------------------
# 日期识别辅助
# ------------------------------------------------------------------
NRIC_PATTERN = re.compile(r"\d{6}\s*-?\s*\d{2}\s*-?\s*\d{4}")

def get_excluded_spans(text):
    return [m.span() for m in NRIC_PATTERN.finditer(text)]

def overlaps(start, end, spans):
    return any(not (end <= s or start >= e) for s, e in spans)

def is_part_of_longer_number(text, start, end):
    before = text[start - 1] if start > 0 else ""
    after = text[end] if end < len(text) else ""
    return before.isdigit() or after.isdigit()

def clean_and_format_date(raw_date_str):
    if not raw_date_str:
        return ""
    cleaned = re.sub(r"[^a-zA-Z0-9\/\-\.\s]", "", raw_date_str).strip()
    try:
        parsed_dt = parser.parse(cleaned, dayfirst=True)
        if parsed_dt.year < 1990 or parsed_dt.year > 2030:
            return ""
        return parsed_dt.strftime("%Y-%m-%d")
    except Exception:
        return ""

# ------------------------------------------------------------------
# 2. 文档内容读取函数
# ------------------------------------------------------------------
def get_text_from_excel(file_bytes):
    try:
        xls = pd.ExcelFile(io.BytesIO(file_bytes))
        full_text = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            text = df.to_string(index=False, na_rep="") 
            full_text.append(text)
        return "\n".join(full_text)
    except Exception as e:
        print(f"⚠️ Excel 读取警告: {e}")
        return ""

def get_text_from_image(file_bytes):
    try:
        img = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(img, lang='eng')
        return text
    except Exception as e:
        print(f"⚠️ 图片 OCR 警告: {e}")
        return ""

def get_text_from_docx(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"⚠️ Word 读取警告: {e}")
        return ""

def get_text_from_pdf(file_bytes):
    try:
        images = convert_from_bytes(file_bytes, first_page=1, last_page=1, poppler_path=POPPLER_PATH)
        if images:
            return pytesseract.image_to_string(images[0], lang='eng')
        return ""
    except Exception as e:
        print(f"⚠️ PDF OCR 警告: {e}")
        return ""

# ------------------------------------------------------------------
# 3. LM Studio 智能兜底分析
# ------------------------------------------------------------------
def analyze_with_lm_studio(text: str) -> dict:
    try:
        prompt = f"""
        Extract the following fields from the document text in JSON format:
        - Document Type
        - Trade Name
        - Lot No
        - Company Name
        - Document Date

        Text:
        {text[:2000]}
        """
        response = client.chat.completions.create(
            model="openai/gpt-oss-20b",
            messages=[
                {"role": "system", "content": "You are a professional document analysis assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1
        )
        content = response.choices[0].message.content.strip()
        content = re.sub(r"^```json\s*|\s*```$", "", content, flags=re.MULTILINE).strip()
        import json
        return json.loads(content)
    except Exception as e:
        print(f"LLM 文本分析失败: {e}")
        return {}

def extract_handwritten_date_with_llm(file_bytes:bytes, ext: str) -> str:
    try:
        target_bytes = file_bytes
        if ext == '.pdf':
            images = convert_from_bytes(file_bytes, first_page=1, last_page=1, poppler_path=POPPLER_PATH)
            if images:
                img_byte_arr = io.BytesIO()
                images[0].save(img_byte_arr, format='JPEG')
                target_bytes = img_byte_arr.getvalue()
            else:
                return ""

        base64_image = base64.b64encode(target_bytes).decode('utf-8')
        
        prompt = (
            "This document/image contains a handwritten date (such as Issue Date, Expiry Date, or Document Date). "
            "Please read the handwritten text, find the document date, and output ONLY the date formatted as YYYY-MM-DD. "
            "If no date is found, reply with 'NONE'."
        )

        response = client.chat.completions.create(
            model="openai/gpt-oss-20b",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
                        }
                    ]
                }
            ],
            temperature=0.1
        )
        
        result = response.choices[0].message.content.strip()
        match = re.search(r"\d{4}-\d{2}-\d{2}", result)
        if match:
            return match.group(0)
        return ""
    except Exception as e:
        print(f"⚠️ 手写日期 LLM 识别失败: {e}")
        return ""

# ------------------------------------------------------------------
# 4. 正则提取逻辑
# ------------------------------------------------------------------
def extract_docuware_fields(full_text, filename=""):
    data = {"Document Type": DOC_TYPE_MAP["DEFAULT"], "Trade Name": "", "Lot / Push Cart No.": "", "Company Name": "", "Document Date": ""}
    text_upper = full_text.upper()
    filename_upper = filename.upper()
    
    # 1. Document Type
    if "HANDOVER" in text_upper or "HANDOVER" in filename_upper: data["Document Type"] = DOC_TYPE_MAP["HANDOVER"]
    elif "RENEWAL OF TENANCY" in text_upper or "TENANCY RENEWAL" in text_upper or "RENEWAL" in filename_upper: data["Document Type"] = DOC_TYPE_MAP["RENEWAL"]
    elif "TENANCY AGREEMENT" in text_upper or " T.A " in text_upper or " TA " in text_upper or ("TA" in filename_upper and "20" in filename_upper): data["Document Type"] = DOC_TYPE_MAP["TA"]
    elif "LETTER OF OFFER" in text_upper: data["Document Type"] = DOC_TYPE_MAP["LO"]
    elif "DEED" in text_upper and "ASSIGNMENT" in text_upper: data["Document Type"] = DOC_TYPE_MAP["DEED"]
    elif "COVER LETTER" in text_upper or "ENCLOSED HEREWITH" in text_upper or "COVER LETTER" in filename_upper: data["Document Type"] = DOC_TYPE_MAP["COVER"]

    # 2. Lot No
    demised_match = re.search(r"DEMISED PREMISES\s*[:]?\s*(.*?)(?:\n|,)", full_text, re.IGNORECASE)
    if demised_match:
        lot_only = re.search(r"([A-Z0-9\-\.& ]+)", demised_match.group(1))
        if lot_only:
            val = lot_only.group(1).strip()
            if not ("2-18" in val.upper() and "OFFICE" in text_upper):
                data["Lot / Push Cart No."] = val

    if not data["Lot / Push Cart No."]:
        bracket_match = re.search(r"\(\s*([0-9A-Z\-\.& ]{2,30})\s*\)", full_text)
        if bracket_match:
            b_val = bracket_match.group(1).strip()
            if (any(c.isdigit() for c in b_val) or "-" in b_val) and "OFFICE" not in b_val.upper():
                data["Lot / Push Cart No."] = f"LOT NO. {b_val}"

    if not data["Lot / Push Cart No."]:
        lot_matches = re.finditer(r"(?:LOT|UNIT|PREMISE)(?:\s+NO\.?)?\s*[:]?\s*([A-Z0-9\-\.& ]+?)(?=\s{2,}|\n|$)", full_text, re.IGNORECASE)
        for match in lot_matches:
            val = match.group(1).strip()
            if "2-18" in val and "OFFICE" in text_upper: continue
            if re.search(r"^\d{1,2}-(?:JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)", val, re.IGNORECASE): continue
            if len(val) < 2: continue
            if not re.search(r"\d", val): continue
            data["Lot / Push Cart No."] = val
            break

    # 3. Company Name
    tenant_match = re.search(r"TENANT\s*[:]\s*(.*?)(?:\n|$)", full_text, re.IGNORECASE)
    if tenant_match: data["Company Name"] = tenant_match.group(1).strip()
    else:
        re_match = re.search(r"RE\s*[:]\s*(.*SDN\.?\s*BHD\.?)", full_text, re.IGNORECASE)
        if re_match: data["Company Name"] = re_match.group(1).strip()
        else:
            sdn_matches = re.findall(r"(.*SDN\.?\s*BHD\.?)", full_text, re.IGNORECASE)
            for name in sdn_matches:
                clean_name = name.strip()
                upper_name = clean_name.upper()
                if "SUNWAY" in upper_name or "RHB TRUSTEES" in upper_name or "MANAGEMENT" in upper_name: continue
                data["Company Name"] = clean_name
                break

    # 4. Trade Name
    trade_match = re.search(r"(?:TRADE NAME|TRADING AS|OUTLET NAME)\s*[:]\s*(.*?)(?:\n|$)", full_text, re.IGNORECASE)
    if trade_match: data["Trade Name"] = trade_match.group(1).strip()
    else:
        data["Trade Name"] = data["Company Name"]
        if data["Trade Name"] and "SDN" in data["Trade Name"].upper():
            data["Trade Name"] = re.sub(r"\s+SDN\.?\s*BHD\.?.*", "", data["Trade Name"], flags=re.IGNORECASE).strip()

    # 5. Date
    excluded_spans = get_excluded_spans(full_text)
    raw_date = ""

    for m in re.finditer(r"(\d{1,2}.{0,3}\s+(?:JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)[a-z]*\s+20\d{2})", full_text, re.IGNORECASE):
        if overlaps(m.start(), m.end(), excluded_spans): continue
        if is_part_of_longer_number(full_text, m.start(), m.end()): continue
        raw_date = m.group(1)
        break

    if not raw_date:
        for m in re.finditer(r"(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{4})", full_text):
            if overlaps(m.start(), m.end(), excluded_spans): continue
            if is_part_of_longer_number(full_text, m.start(), m.end()): continue
            raw_date = m.group(1)
            break

    data["Document Date"] = clean_and_format_date(raw_date)
    return data

# ------------------------------------------------------------------
# 5. FastAPI 接口与后台处理
# ------------------------------------------------------------------
def process_file(file_bytes: bytes, filename: str) -> dict:
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    try:
        if ext == '.pdf':
            text = get_text_from_pdf(file_bytes)
        elif ext == '.docx':
            text = get_text_from_docx(file_bytes)
        elif ext in ['.xlsx', '.xls']:
            text = get_text_from_excel(file_bytes)
        elif ext in ['.jpg', '.jpeg', '.png']:
            text = get_text_from_image(file_bytes)

        data = extract_docuware_fields(text, filename)

        if not data["Lot / Push Cart No."] or not data["Company Name"] or not data["Document Date"]:
            print(f"🤖 正在为 {filename} 触发 LLM 智能兜底...")
            llm_res = analyze_with_lm_studio(text)

            if not data["Lot / Push Cart No."] and llm_res.get("Lot No"):
                data["Lot / Push Cart No."] = llm_res.get("Lot No")
            if not data["Company Name"] and llm_res.get("Company Name"):
                data["Company Name"] = llm_res.get("Company Name")
            if not data["Document Date"] and llm_res.get("Document Date"):
                data["Document Date"] = clean_and_format_date(llm_res.get("Document Date"))

        if not data["Document Date"]:
            handwritten_date = extract_handwritten_date_with_llm(file_bytes, ext)
            if handwritten_date:
                data["Document Date"] = handwritten_date

        return {
            'File Name': filename,
            'Document Type': data['Document Type'],
            'Trade Name': data['Trade Name'],
            'Lot No': data['Lot / Push Cart No.'],
            'Company Name': data['Company Name'],
            'Document Date': data['Document Date']
        }
    except Exception as e:
        print(f"❌ 深度处理失败: {filename} -> {e}")
        return {
            'File Name': filename,
            'Document Type': 'ERROR',
            'Trade Name': '',
            'Lot No': '',
            'Company Name': f'Error: {str(e)}',
            'Document Date': ''
        }


async def run_analysis_job(job_id: str, files_data: list):
    """
    背景任务：逐个处理文件，增加取消指令检查。

    关键改动：process_file 内部包含同步阻塞调用（OCR、pdf2image、以及对
    LM Studio 的同步 HTTP 请求），如果直接在这个 async 函数里调用会卡住
    整个事件循环，导致其他请求（包括 /cancel、/progress、甚至新的 /start）
    在这段时间内完全无法响应。用 asyncio.to_thread 把它丢进线程池执行，
    避免阻塞事件循环，这样取消按钮和轮询才能正常工作。
    """
    output_rows = []

    for i, (filename, file_bytes) in enumerate(files_data):
        # 🛑 在处理每一个文件前检查该 job_id 是否已被取消
        if job_id in CANCELLED_JOBS:
            print(f"🛑 Job {job_id} 已被用户手动停止！直接退出后台处理...")
            if job_id in JOBS:
                JOBS[job_id]["status"] = "cancelled"
            CANCELLED_JOBS.discard(job_id)  # 清理集合（discard 不存在时不会报错）
            return

        # 关键改动：用线程池跑同步阻塞代码，不卡住事件循环
        row = await asyncio.to_thread(process_file, file_bytes, filename)
        output_rows.append(row)
        JOBS[job_id]["done"] = i + 1

    # 循环结束后再检查一次，防止最后一个文件处理期间被取消
    if job_id in CANCELLED_JOBS:
        print(f"🛑 Job {job_id} 在收尾阶段被取消！")
        if job_id in JOBS:
            JOBS[job_id]["status"] = "cancelled"
        CANCELLED_JOBS.discard(job_id)
        return

    # 生成 CSV
    output = io.StringIO()
    fieldnames = ['File Name', 'Document Type', 'Trade Name', 'Lot No', 'Company Name', 'Document Date']
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    writer.writerows(output_rows)

    JOBS[job_id]["result"] = output.getvalue().encode('utf-8-sig')
    JOBS[job_id]["status"] = "finished"


@app.post("/api/analyze-folder/start")
async def start_analysis(files: list[UploadFile] = File(...)):
    job_id = str(uuid.uuid4())

    files_data = []
    for file in files:
        content = await file.read()
        files_data.append((file.filename, content))

    JOBS[job_id] = {"total": len(files_data), "done": 0, "status": "processing", "result": None}

    asyncio.create_task(run_analysis_job(job_id, files_data))

    return {"job_id": job_id, "total": len(files_data)}


# 🛑 取消任务接口
@app.post("/api/analyze-folder/cancel/{job_id}")
async def cancel_analysis(job_id: str):
    if job_id in JOBS:
        CANCELLED_JOBS.add(job_id)
        JOBS[job_id]["status"] = "cancelled"
        return {"status": "cancelled", "job_id": job_id}
    raise HTTPException(status_code=404, detail="Job not found")


@app.get("/api/analyze-folder/progress/{job_id}")
async def get_progress(job_id: str):
    job = JOBS.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    percent = int((job["done"] / job["total"]) * 100) if job["total"] > 0 else 0
    return {"done": job["done"], "total": job["total"], "percent": percent, "status": job["status"]}


@app.get("/api/analyze-folder/download/{job_id}")
async def download_result(job_id: str):
    job = JOBS.get(job_id)
    if not job or job["status"] != "finished":
        raise HTTPException(status_code=400, detail="Job not finished or was cancelled")

    return StreamingResponse(
        io.BytesIO(job["result"]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=docuware_export.csv"}
    )